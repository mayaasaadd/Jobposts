import streamlit as st
from langchain_community.chat_models import ChatOllama
from langchain.prompts import PromptTemplate
import pdfplumber
from docx import Document
from io import BytesIO
from docx.shared import Pt
from docx.shared import RGBColor




# --- Page Config ---
st.set_page_config(
    page_title="Job Description → Job Post Generator",
    layout="centered",
    page_icon="📄"
)

# --- Custom CSS for Branding ---
st.markdown("""
<style>
/* Main page buttons - Ahlyia red */
.stButton>button {
    background-color: #C8102E;  
    color: white;
    font-weight: 600;
    border-radius: 6px;
    padding: 0.5rem 1rem;
    font-size: 1rem;
}

/* Hover effect */
.stButton>button:hover {
    background-color: #E03C4E; 
}

/* File uploader browse button */
.css-1aumxhk button {
    background-color: #C8102E !important;
    color: white !important;
    font-weight: 600 !important;
    border-radius: 6px !important;
}
.css-1aumxhk button:hover {
    background-color: #E03C4E !important;
}

/* Sidebar background and buttons */
[data-testid="stSidebar"] {
    background-color: #003366;
    color: white;
}
[data-testid="stSidebar"] h1,
[data-testid="stSidebar"] h2,
[data-testid="stSidebar"] h3,
[data-testid="stSidebar"] p,
[data-testid="stSidebar"] span,
[data-testid="stSidebar"] label {
    color: white;
}
[data-testid="stSidebar"] button {
    background-color: #C8102E;
    color: white;
    font-weight: 600;
    border-radius: 6px;
}
[data-testid="stSidebar"] button:hover {
    background-color: #E03C4E;
}

/* Preview containers and text areas */
.job-preview, .stTextArea>div>div>textarea {
    color: #111111 !important;
}
.css-1cpxqw2 input {
    color: #111111 !important;
}
.css-1cpxqw2 input::placeholder {
    color: #555555 !important;
}

/* Individual job post buttons (dark blue) */
.individual-download button {
    background-color: #003366 !important;
    color: white !important;
    font-weight: 600 !important;
    border-radius: 6px !important;
}
.individual-download button:hover {
    background-color: #0056A0 !important;
}

            
/* Sidebar slider - red theme */
[data-testid="stSidebar"] .stSlider {
    accent-color: #C8102E;  /* Modern browsers respect this for track & knob */
}



</style>
            
""", unsafe_allow_html=True)

# --- Centered Logo ---
col1, col2, col3 = st.columns([1, 2, 1])
with col2:
    st.image("Unknown.png", width=400)
st.markdown("<br>", unsafe_allow_html=True)

# --- Main Banner ---
st.markdown("""
<div style="
    background-color: #003366;
    padding: 30px 25px;
    border-radius: 12px;
    box-shadow: 0 6px 10px rgba(0,0,0,0.15);
    text-align: center;
    color: white;
    margin-bottom: 25px;
    position: relative;
">
    <h1 style="margin:0; font-size:2rem; font-weight:700;"> Job Post Generator</h1>
    <div style="
        width: 80px;
        height: 4px;
        background-color: #C8102E;
        margin: 10px auto 0;
        border-radius: 2px;
    "></div>
            <h2 style="margin:0; font-size:1.5rem;">📂 Upload Job Descriptions</h2>
    <p style="margin:5px 0 15px; font-size:1rem;">
        Upload your PDF or DOCX Job Descriptions. The system will generate professional LinkedIn-ready job posts for El Ahlyia Healthcare.
    </p>
            
</div>
""", unsafe_allow_html=True)


# --- Session State ---
if "editable_posts" not in st.session_state:
    st.session_state.editable_posts = []  # Initialize the list to avoid errors
existing_files = {p["filename"] for p in st.session_state.editable_posts}
if "search_filter" not in st.session_state:
    st.session_state.search_filter = ""


# --- Company Description ---
COMPANY_DESCRIPTION = (
    "Company Description:\n"
    "El Ahlyia Healthcare (EHC) is a leading medical supplies distributor serving the Egyptian market "
    "with advanced medical technologies since 1995. We are dedicated to improving healthcare outcomes "
    "and are recognized for delivering reliable and innovative medical solutions."
)


# --- Helper Functions ---
def extract_text(uploaded_file):
    jd_text = ""
    if uploaded_file.type == "application/pdf":
        with pdfplumber.open(uploaded_file) as pdf:
            for page in pdf.pages:
                jd_text += page.extract_text() or ""
    else:
        doc = Document(uploaded_file)
        for para in doc.paragraphs:
            jd_text += para.text + " "
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    jd_text += cell.text + " "
    return jd_text.strip()

def format_bullets(text):
    """
    Clean and format bullets from LLM output:
    - Removes intro lines like "Here is the rewritten job post"
    - Removes extra markdown (** or *)
    - Converts lines starting with • or - into uniform format
    """
    lines = []
    for line in text.split("\n"):
        line = line.strip()
        if not line:
            continue
        # Remove LLM intro
        if "Here is the rewritten job post" in line or line == "•":
            continue
        # Remove markdown symbols
        line = line.replace("**", "").replace("*", "").strip()
        # Convert bullets to -
        if line.startswith(("•", "-")):
            line = line.lstrip("•- ").strip()
            lines.append(f"- {line}")
        else:
            lines.append(line)
    return "\n".join(lines)


def create_docx(content, default_title):
    from docx import Document
    from io import BytesIO

    doc = Document()

    # Split content into lines
    lines = content.split("\n")

    # --- CLEAN LINES ---
    cleaned_lines = []
    for line in lines:
        line = line.strip()
        if not line:
            continue
        # Remove LLM intro and stray bullets
        if "Here's the rewritten job post" in line or line == "•":
            continue
        # Remove markdown symbols
        line = line.replace("**", "").replace("*", "").strip()
        cleaned_lines.append(line)
    
    lines = cleaned_lines

    # --- EXTRACT JOB TITLE ---
    job_title = default_title
    if lines and not any(lines[0].startswith(section) for section in [
        "Company Description", "Role Description", "Qualifications",
        "Job Requirements", "Reasons to Join"
    ]):
        job_title = lines[0]
        lines = lines[1:]

    # Add job title as main heading
    doc.add_heading(job_title, level=1)

    # --- ADD REST OF CONTENT ---
    for line in lines:
        line = line.strip()
        if not line:
            continue

        # Section headings
        if any(line.startswith(section) for section in [
            "Company Description", "Role Description", "Qualifications",
            "Job Requirements", "Reasons to Join"
        ]):
            p = doc.add_paragraph()
            p.add_run(line).bold = True

        # Bullet points starting with - or •
        elif line.startswith(("•", "-")):
            doc.add_paragraph(line[1:].strip(), style="List Bullet")

        # Lines in "Reasons to Join" or other items without bullets
        elif any(keyword in line for keyword in [
            "Be part of", "Work on", "Collaborate", "Enjoy", "Contribute"
        ]):
            doc.add_paragraph(line, style="List Bullet")

        # Normal paragraph
        else:
            doc.add_paragraph(line)

    # Save to BytesIO
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer







# --- Sidebar ---

with st.sidebar:
    st.image("Unknown.png", width=300)
    st.title(" ⚙ Settings")
    model_name = st.selectbox("Select LLM Model", ["llama3:latest"], index=0)
    temp = st.slider("Creativity (Temperature)", 0.0, 1.0, 0.7)
    if st.button(" 🗑 Clear All Drafts", use_container_width=True):
        st.session_state.editable_posts = []
        st.session_state.search_filter = ""
        st.rerun()



# --- File Upload Section ---
uploaded_files = st.file_uploader(
    "", type=["pdf", "docx"], accept_multiple_files=True
)


# --- Generate Job Posts ---
if st.button("Generate Job Posts"):
    if not uploaded_files:
        st.warning("Please upload at least one Job Description file to generate drafts.")
    else:
        chat = ChatOllama(model=model_name, temperature=temp, base_url="https://api.ollama.com")
        prompt_template = PromptTemplate(
            input_variables=["jd_text"],
            template=(
                "You are an HR Talent Acquisition professional.\n"
                "Rewrite the following Job Description into a concise, professional LinkedIn-style job post.\n"
                "Use hyphens '-' for bullet points.\n"
                "Maintain the following sections:Job Title, Role Description, Qualifications, Job Requirements, Reasons to Join El Ahlyia.\n"
                "Make Sure Bullet Points are Found in Qualifications and Job Requirements"
                "Job Description Content:\n{jd_text}"
            ),
        )

        existing_files = {p["filename"] for p in st.session_state.editable_posts}

        # Containers for progress
        progress_container = st.container()
        overall_progress = st.progress(0)

        for idx, uploaded_file in enumerate(uploaded_files, start=1):
            if uploaded_file.name in existing_files:
              continue
            with progress_container:
                # Yellow bar - processing
                st.markdown(
                    f"""
                    <div style="
                        background-color: #FFF3CD;
                        padding: 10px;
                        border-radius: 5px;
                        border: 1px solid #FFE69C;
                        color: #856404;
                        margin-bottom:5px;">
                        Processing: {uploaded_file.name}
                    </div>
                    """,
                    unsafe_allow_html=True
                )

            jd_text = extract_text(uploaded_file)
            if not jd_text:
                jd_text = "Job Description content could not be extracted from the file."

            try:
              raw_post = chat.predict(prompt_template.format(jd_text=jd_text))
              job_post = format_bullets(raw_post)
              final_post = f"{COMPANY_DESCRIPTION}\n\n{job_post}"
            except Exception as e:
             final_post = f"{COMPANY_DESCRIPTION}\n\n⚠️ Failed to generate job post: {str(e)}"

            st.session_state.editable_posts.append({
             "filename": uploaded_file.name,
             "content": final_post
            })


            with progress_container:
                # Green bar - completed
                st.markdown(
                    f"""
                    <div style="
                        background-color: #D4EDDA;
                        padding: 10px;
                        border-radius: 5px;
                        border: 1px solid #C3E6CB;
                        color: #155724;
                        margin-bottom:5px;">
                        Completed: {uploaded_file.name}
                    </div>
                    """,
                    unsafe_allow_html=True
                )
            overall_progress.progress(idx / len(uploaded_files))

        st.success(f"Generated {len(st.session_state.editable_posts)} draft job post(s). Edit and download below.")


# --- Search / Filter ---
if st.session_state.get("editable_posts"):
    st.subheader(" 🔍 Search / Filter Job Posts")
    st.session_state.search_filter = st.text_input(
        "Enter job title or keyword to filter drafts", st.session_state.get("search_filter", "")
    )
    search_query = st.session_state.search_filter.lower()
    filtered_posts = []
    for idx, post in enumerate(st.session_state.editable_posts):
        key = f"final_{idx+1}"
        content_to_search = st.session_state.get(key, post["content"])
        first_lines = "\n".join(content_to_search.split("\n")[:5]).lower()
        if search_query in post["filename"].lower() or search_query in first_lines:
            filtered_posts.append((idx, post))
    if not filtered_posts:
        st.info("No job posts match the search filter.")


# --- Edit & Download Individual Posts ---
if st.session_state.get("editable_posts") and filtered_posts:
    st.header(" ✏ Edit & Download Individual Job Posts")
    for idx, post in filtered_posts:
        with st.expander(f"Job Post {idx+1} — {post['filename']}", expanded=False):
            final_text = st.text_area(
                f"Edit Job Post {idx+1}", value=st.session_state.get(f"final_{idx+1}", post["content"]),
                height=300, key=f"final_{post['filename']}"

            )
            # Preview container
            st.markdown(
                f'<div class="job-preview" style="background-color:#f7f7f7; padding:20px; border-radius:10px; border:1px solid #ddd; white-space:pre-wrap; margin-top:10px;">{final_text}</div>',
                unsafe_allow_html=True
            )
            # Individual download button (dark blue)
            buffer = create_docx(final_text, f"Job Post {idx+1} — {post['filename']}")
            st.download_button(
                label=f"Download Job Post {idx+1} (DOCX)",
                data=buffer,
                file_name=f"{post['filename'].replace('.','_')}_JobPost.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                key=f"download_{idx}",
                help="Download this job post",
                use_container_width=True
            )


# --- Batch Download ---
if st.button("Download All Drafts (DOCX)"):
    master_doc = Document()

    # --- Header for Master Document ---
    header_title = master_doc.add_paragraph("LinkedIn Job Posts Generator")
    run = header_title.runs[0]
    run.bold = True
    run.font.size = Pt(24)
    header_title.alignment = 1  # center

    header_subtitle = master_doc.add_paragraph("El Ahlyia Healthcare — Generated Job Posts")
    run = header_subtitle.runs[0]
    run.italic = True
    run.font.color.rgb = RGBColor(0, 153, 255)  # Same Dark Blue
    run.font.size = Pt(14)
    header_subtitle.alignment = 1  # center

    master_doc.add_paragraph("")  # small spacing after header

    # Loop through all posts
    for idx, post in enumerate(st.session_state.editable_posts, start=1):
        final_text = st.session_state.get(f"final_{idx}", post["content"])

        # Add heading for each job post
        job_post_title = f"Job Post {idx} — {post['filename']}"
        master_doc.add_heading(job_post_title, level=1)

        # Split content into lines
        lines = final_text.split("\n")
        lines = [line.strip() for line in lines if line.strip()]

        # Add each line with formatting
        for line in lines:
            # Section headings
            if any(line.startswith(section) for section in [
                "Company Description", "Role Description", "Qualifications",
                "Job Requirements", "Reasons to Join"
            ]):
                p = master_doc.add_paragraph()
                p.add_run(line).bold = True

            # Bullet points
            elif line.startswith(("-", "•")):
                master_doc.add_paragraph(line.lstrip("-• ").strip(), style="List Bullet")

            # Regular paragraph
            else:
                master_doc.add_paragraph(line)

        master_doc.add_paragraph("")  # spacing between posts

    # Save to BytesIO
    buffer = BytesIO()
    master_doc.save(buffer)
    buffer.seek(0)

    st.download_button(
        label="Download All Job Posts (DOCX)",
        data=buffer,
        file_name="All_Job_Posts.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        use_container_width=True
    )



